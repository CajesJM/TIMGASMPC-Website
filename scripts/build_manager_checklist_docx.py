from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "TIMGAS_MPC_Manager_Data_Checklist.docx"

GREEN = "123F2F"
DARK_GREEN = "0B2D22"
MID_GREEN = "23694F"
LIGHT_GREEN = "E3EFE9"
GOLD = "C18C27"
LIGHT_GOLD = "F6EBD3"
CHARCOAL = "24332D"
MUTED = "5E6D67"
WHITE = "FFFFFF"
LINE = "DCE4DF"


SECTIONS = [
    ("1. Official cooperative information", [
        "Complete registered name and preferred public name",
        "Meaning of TIMGAS, if it should be published",
        "Founding or registration date and approved CDA registration details",
        "Short official history or background",
        "Mission, vision, and core values",
        "Official tagline or slogan",
        "Official logo in SVG, PNG, or high-resolution format",
        "Approved brand colors",
    ]),
    ("2. Contact and office details", [
        "Complete office address",
        "Google Maps link or map coordinates",
        "Official mobile and landline numbers",
        "Official public email address",
        "Manager email address for application notifications",
        "Office hours and closure days",
        "Official Facebook and other social-media links",
    ]),
    ("3. Verified public statistics", [
        "Number of active members",
        "Number of communities or barangays served",
        "Number of active services and programs",
        "Number of farmers, people, or businesses assisted",
        "Total loans or assistance released, if approved for publication",
        "Important awards, certifications, and achievements",
        "Date each published figure was last verified",
    ]),
    ("4. Membership requirements", [
        "Eligibility, age, and location requirements",
        "Regular and associate membership rules",
        "Required documents and accepted government IDs",
        "Application fee and payment methods",
        "Initial share capital, savings deposit, and other required fees",
        "Membership orientation requirements",
        "Expected processing time",
        "Approval, rejection, withdrawal, and termination rules",
        "Official membership application form",
        "Member rights and responsibilities",
    ]),
    ("5. Online application process", [
        "Application types the website should accept",
        "Required and optional applicant fields",
        "Required uploaded documents for each application type",
        "Allowed file formats, maximum file size, and maximum file count",
        "Person responsible for reviewing and approving applications",
        "Approved status workflow: New, In Review, Incomplete, For Verification, Approved, Rejected, Withdrawn, or Archived",
        "Reasons an application may be rejected",
        "Procedure for incomplete or incorrect submissions",
        "Whether applicants may correct a submitted application",
        "Application and document retention period",
    ]),
    ("6. Services and programs", [
        "Official name and public description of every service",
        "Eligibility requirements",
        "Required documents",
        "Fees, interest rates, repayment terms, and penalties when applicable",
        "Application procedure and expected processing time",
        "Responsible contact person",
        "Whether each service is currently active",
        "Related downloadable forms",
    ]),
    ("7. Officers and management", [
        "Full name and official position of every person to be displayed",
        "Term start and end dates",
        "Short professional biography",
        "Official photograph",
        "Written permission to publish the name, photograph, and profile",
        "Correct website display order",
    ]),
    ("8. Website content", [
        "Current announcements, advisories, events, and deadlines",
        "Training schedules and general assembly information",
        "Official downloadable forms, reports, and policies",
        "Real member testimonials with publication consent",
        "Cooperative photographs with captions and publication permission",
        "Person authorized to review, approve, and publish website content",
    ]),
    ("9. Privacy and legal requirements", [
        "Approved Privacy Notice",
        "Applicant consent statement",
        "Website Terms of Use",
        "Data-retention and deletion policy",
        "Data Protection Officer or privacy contact",
        "Procedure for correcting applicant information and handling complaints",
        "Permission wording for photographs and testimonials",
        "Confirmation that every collected field and document is necessary",
    ]),
    ("10. Manager account and notifications", [
        "Manager full name, official position, and work email",
        "Recovery mobile number and multi-factor authentication preference",
        "Backup administrator and manager-replacement procedure",
        "Required dashboard reports, searches, filters, printing, and exports",
        "Application actions the manager is authorized to perform",
        "Manager notification email and backup recipient",
        "Approved confirmation, approval, rejection, and incomplete-application email wording",
        "Approved sender name and reply-to email address",
    ]),
    ("11. Technical ownership", [
        "Google account that will own the Firebase project",
        "Official website domain and domain owner",
        "Preferred Firebase project name",
        "Expected number of applications per month",
        "Backup frequency and responsible person",
        "Person authorized to approve deployment",
    ]),
]

PLACEHOLDERS = [
    "Barangay Timogas office address",
    "Phone number and hello@timgasmpc.org email address",
    "Founding year: 2008",
    "Active members: 1,200+",
    "Member loans supported: PHP 32M+",
    "Six listed cooperative services",
    "Current announcements and dates",
    "Sample member testimonial",
    "Generated farmer hero image",
    "Growing stronger, together tagline",
    "Sample manager and application records in the dashboard preview",
]


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        borders.append(left)
        p_pr.append(borders)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def set_page_furniture(section):
    for footer_part in (section.footer, section.even_page_footer, section.first_page_footer):
        footer = footer_part.paragraphs[0]
        footer.paragraph_format.space_before = Pt(0)
        add_page_field(footer)


def add_checkbox_numbering(document):
    numbering = document.part.numbering_part.element
    existing = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=-1) + 1
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "☐")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Segoe UI Symbol")
    r_fonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MID_GREEN)
    r_pr.append(color)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_checkbox(document, text, num_id):
    paragraph = document.add_paragraph(style="Checklist Item")
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_ref])
    paragraph._p.get_or_add_pPr().insert(0, num_pr)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=CHARCOAL)
    return paragraph


def build_document():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    doc.core_properties.title = "TIMGAS MPC Important Manager Data Checklist"
    doc.core_properties.subject = "Essential information required for the website and Firebase backend"
    doc.core_properties.author = "TIMGAS Multi-Purpose Cooperative"
    doc.core_properties.keywords = "TIMGAS, cooperative, website, manager checklist"

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, GREEN, 18, 10),
        ("Heading 2", 13, GREEN, 14, 7),
        ("Heading 3", 12, DARK_GREEN, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    checklist_style = styles.add_style("Checklist Item", WD_STYLE_TYPE.PARAGRAPH)
    checklist_style.base_style = normal
    checklist_style.paragraph_format.space_after = Pt(4)
    checklist_style.paragraph_format.line_spacing = 1.25
    checklist_style.paragraph_format.widow_control = True

    set_page_furniture(section)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    r = kicker.add_run("MANAGER INFORMATION PACKET")
    set_run_font(r, size=9, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("Important Data Checklist")
    set_run_font(r, size=28, color=DARK_GREEN, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("Essential information needed to complete the TIMGAS MPC website and prepare the secure Firebase backend.")
    set_run_font(r, size=11.5, color=MUTED)

    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Inches(0.16)
    note.paragraph_format.right_indent = Inches(0.10)
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(14)
    note.paragraph_format.line_spacing = 1.15
    shade_paragraph(note, LIGHT_GOLD, GOLD)
    r = note.add_run("Important: ")
    set_run_font(r, size=9.5, color=DARK_GREEN, bold=True)
    r = note.add_run("Do not enter passwords, Gmail passwords, API keys, service-account files, or other production secrets in this document.")
    set_run_font(r, size=9.5, color=CHARCOAL)

    details = doc.add_paragraph()
    details.paragraph_format.space_after = Pt(2)
    r = details.add_run("Manager name: ")
    set_run_font(r, size=10, color=GREEN, bold=True)
    r = details.add_run("________________________________________")
    set_run_font(r, size=10, color=MUTED)
    details2 = doc.add_paragraph()
    details2.paragraph_format.space_after = Pt(14)
    r = details2.add_run("Date started:  ")
    set_run_font(r, size=10, color=GREEN, bold=True)
    r = details2.add_run("____________________")
    set_run_font(r, size=10, color=MUTED)

    num_id = add_checkbox_numbering(doc)
    for heading, items in SECTIONS:
        doc.add_paragraph(heading, style="Heading 1")
        for item in items:
            add_checkbox(doc, item, num_id)

    doc.add_paragraph("Website placeholders to confirm or replace", style="Heading 1")
    intro = doc.add_paragraph("The following values are currently used only as draft content and must be verified before launch:")
    intro.paragraph_format.keep_with_next = True
    for item in PLACEHOLDERS:
        add_checkbox(doc, item, num_id)

    doc.add_paragraph("Final approval", style="Heading 1")
    for item in (
        "All public information has been reviewed and is accurate",
        "Financial figures and service terms have been approved",
        "People shown on the website have provided publication consent",
        "Application fields and uploaded documents are necessary and approved",
        "Privacy and consent documents have been reviewed",
        "All placeholders have been replaced or explicitly approved",
        "The manager authorizes the approved website content for publication",
    ):
        add_checkbox(doc, item, num_id)

    doc.add_paragraph()
    for label in ("Manager name", "Signature", "Date reviewed"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, color=GREEN, bold=True)
        r = p.add_run("________________________________________")
        set_run_font(r, size=10.5, color=MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
